# Andrea, Grace, Jack, Peter
# CRT Scraper
# Task 1: 3 weeks of dates for searching

# Objective: scrape 3 weeks (sept. 15 -22 2020, sept 22 - 29 2020, and august 24-31 2022)
# for one news site. Port data and full text to a combined document (docx or pdf) for viewing.

import docx
#from docx import Document
from pygooglenews import GoogleNews
from newspaper import Article
from newspaper import Config
import requests
from scrapingbee import ScrapingBeeClient
from urllib3.util.retry import Retry
from requests.adapters import HTTPAdapter
import re

#config will allow us to access the specified url for which we are #not authorized. Sometimes we may get 403 client error while parsing #the link to download the article.

count = 0
count2 = 0
count3 = 0
our_api_key = "O8IZ2YUTDJ0TK9YV2LY7EAFMM0TB9ID6PGPVZOYA7JIGK33RILWEG28T2VRNBCVRVTV3499R9Z1OZ580"
gn = GoogleNews()

# gn.search(query: str, helper = True, when = None, from_ = None, to_ = None, proxies=None, scraping_bee=None)
wapo_articles1 = gn.search('Critical Race Theory site:washingtonpost.com before:2020-09-22 after:2020-09-15')
for article in wapo_articles1["entries"]:
    count+= 1
    #print(article["link"])
print(count)

# print("End of week")

wapo_articles2 = gn.search('Critical Race Theory site:washingtonpost.com before:2020-09-29 after:2020-09-21')
for article in wapo_articles2["entries"]:
    count2+= 1

print(count2)

# print("End of week")

wapo_articles3 = gn.search('Critical Race Theory site:washingtonpost.com before:2022-09-01 after:2022-08-24')
for article in wapo_articles3["entries"]:
    count3+= 1
print(count3)

# /Users/peterbutler/Developer/SEMINAR/week1.docx 
# /Users/peterbutler/Developer/SEMINAR/week2.docx 
# /Users/peterbutler/Developer/SEMINAR/week3.docx 


wapo_doc_1 = docx.Document()
wapo_doc_2 = docx.Document()
wapo_doc_3 = docx.Document()


client = ScrapingBeeClient(api_key=our_api_key)



def doc_maker(your_doc, links, save_link):
    if len(links) == 0:
        return
        
    for link in links:
        raw_html = client.get(link["link"])

        article = Article('')
        article.download(raw_html.content)
        article.parse()

        if article.title:
            your_doc.add_heading(article.title)
        if article.authors:
            your_doc.add_paragraph(article.authors)
        if article.publish_date:
            your_doc.add_paragraph(str(article.publish_date))
        if article.text:
            your_doc.add_paragraph(article.text)
        your_doc.add_page_break()

    your_doc.save(save_link)

doc_maker(wapo_doc_1, wapo_articles1["entries"], "/Users/peterbutler/Developer/SEMINAR/week1.docx")
doc_maker(wapo_doc_2, wapo_articles2["entries"], "/Users/peterbutler/Developer/SEMINAR/week2.docx")
doc_maker(wapo_doc_3, wapo_articles3["entries"], "/Users/peterbutler/Developer/SEMINAR/week3.docx")


